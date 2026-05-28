"""
Generate sample rental contract Word documents for fresh university graduates.
"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

output_dir = r"D:\aaa\contract-review-copilot\sample_contracts"
os.makedirs(output_dir, exist_ok=True)

def create_contract(filename, data):
    doc = Document()
    title = doc.add_heading(data["title"], 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"合同编号：{data['contract_no']}")
    doc.add_paragraph("")

    doc.add_heading("一、合同当事人", level=2)
    doc.add_paragraph(f"甲方（出租方）：{data['lessor']}")
    doc.add_paragraph(f"乙方（承租方）：{data['lessee']}")

    doc.add_heading("二、租赁房屋基本情况", level=2)
    doc.add_paragraph(f"房屋地址：{data['property_address']}")
    doc.add_paragraph(f"房屋面积：{data['area']} 平方米")
    doc.add_paragraph(f"户型：{data['layout']}")
    doc.add_paragraph(f"房屋用途：仅限乙方居住使用")
    doc.add_paragraph(f"房屋现状：{data['condition']}")

    doc.add_heading("三、租赁期限", level=2)
    doc.add_paragraph(f"租赁开始日期：{data['start_date']}")
    doc.add_paragraph(f"租赁结束日期：{data['end_date']}")
    doc.add_paragraph(f"租赁期限：{data['duration']}")

    doc.add_heading("四、租金及支付方式", level=2)
    doc.add_paragraph(f"月租金：人民币 {data['monthly_rent']} 元")
    doc.add_paragraph(f"租金支付方式：{data['payment_method']}")
    doc.add_paragraph(f"押金：人民币 {data['deposit']} 元（{data['deposit_months']}个月租金）")
    doc.add_paragraph(f"押金退还条件：{data['deposit_conditions']}")

    doc.add_heading("五、费用约定", level=2)
    doc.add_paragraph(f"水费：{data['water_fee']}")
    doc.add_paragraph(f"电费：{data['electricity_fee']}")
    doc.add_paragraph(f"燃气费：{data['gas_fee']}")
    doc.add_paragraph(f"物业费：{data['property_fee']}")
    doc.add_paragraph(f"网络费：{data['internet_fee']}")
    doc.add_paragraph(f"中介费：{data['agency_fee']}")

    doc.add_heading("六、双方权利与义务", level=2)
    for item in data['obligations']:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading("七、违约责任", level=2)
    for item in data['penalties']:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading("八、其他约定", level=2)
    for item in data['other_terms']:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph("")
    doc.add_paragraph("甲方（签字/盖章）：______________    乙方（签字）：______________")
    doc.add_paragraph(f"签订日期：{data['sign_date']}")
    doc.add_paragraph(f"签订地点：{data['sign_location']}")

    filepath = os.path.join(output_dir, filename)
    doc.save(filepath)
    print(f"Generated: {filepath}")


# ========== 合同1：正规长租公寓（无风险） ==========
create_contract("01_正规长租公寓合同.docx", {
    "title": "房屋租赁合同",
    "contract_no": "ZA-2024-001",
    "lessor": "万科泊寓（深圳）公寓管理有限公司（统一社会信用代码：91440300MA5DD12345）",
    "lessee": "林晓雨（身份证：440301199901011234）",
    "property_address": "深圳市南山区西丽街道留仙洞总部基地万科泊寓·创客社区A栋808室",
    "area": "35",
    "layout": "一室一厅一卫（Loft复式）",
    "condition": "精装修，配备空调、热水器、洗衣机、冰箱、床、衣柜、书桌等家具家电",
    "start_date": "2024年7月1日",
    "end_date": "2025年6月30日",
    "duration": "12个月",
    "monthly_rent": "2,800",
    "payment_method": "每月5日前通过泊寓APP以银行代扣方式支付",
    "deposit": "5,600（2个月租金）",
    "deposit_months": "2",
    "deposit_conditions": "租赁期满或提前解除时，乙方结清全部费用且房屋及设施完好（正常使用磨损除外），甲方在5个工作日内无息退还全部押金",
    "water_fee": "按实际使用量由物业代收，每季度结算",
    "electricity_fee": "按国家标准电价由物业代收，独立电表计量",
    "gas_fee": "无燃气",
    "property_fee": "每月150元（含公共区域保洁、快递代收、安保服务）",
    "internet_fee": "每月60元（已接入100M宽带）",
    "agency_fee": "无（品牌公寓直租）",
    "obligations": [
        "甲方应确保房屋及附属设施处于正常使用状态，乙方报修后24小时内响应",
        "乙方应按时缴纳租金及各项费用",
        "乙方不得留宿他人超过连续3天，或在屋内从事违法活动",
        "乙方提前退租须提前30天书面通知甲方，并配合甲方看房"
    ],
    "penalties": [
        "乙方逾期支付租金，每逾期一日按月租金的0.5%收取滞纳金",
        "乙方提前退租且未履行30天通知义务的，押金不予退还",
        "甲方提前收回房屋的，应提前30天通知乙方，并退还剩余租金及押金"
    ],
    "other_terms": [
        "本合同自乙方支付首月租金及押金之日起生效",
        "乙方入住前需完成物品交接清单签署",
        "合同期满，乙方同等条件下享有优先续租权",
        "争议解决：提交合同签订地人民法院诉讼解决"
    ],
    "sign_date": "2024年6月20日",
    "sign_location": "深圳市南山区"
})

# ========== 合同2：隔断房/合租房风险版 ==========
create_contract("02_隔断间合租风险合同.docx", {
    "title": "房屋租赁合同",
    "contract_no": "HZ-2024-015",
    "lessor": "李建国（身份证：330102197505051122）",
    "lessee": "张浩然（身份证：320281199903025678）",
    "property_address": "杭州市西湖区文三路123号华星时代广场B栋2104室（原始户型为三室一厅）",
    "area": "12（隔断间）",
    "layout": "隔断单间（与其他两室共用厨卫）",
    "condition": "简单装修，配备床、衣柜、空调，无独立卫生间",
    "start_date": "2024年8月1日",
    "end_date": "2025年7月31日",
    "duration": "12个月",
    "monthly_rent": "1,500",
    "payment_method": "每季度第一个月第一天以银行转账方式支付",
    "deposit": "4,500（3个月租金）",
    "deposit_months": "3",
    "deposit_conditions": "合同到期乙方搬离并结清所有费用后，甲方验收房屋无损坏时退还押金。如有损坏，从押金中扣除修复费用，余额退还",
    "water_fee": "每季度按人头分摊（预估每季度200元）",
    "electricity_fee": "按电表度数自行向房东缴纳（电费1.2元/度，含公摊电费）",
    "gas_fee": "与其他租客共用，按人头分摊",
    "property_fee": "每月100元",
    "internet_fee": "每月50元（共享带宽）",
    "agency_fee": "一次性收取800元（已包含在首期费用中）",
    "obligations": [
        "乙方须遵守作息时间，晚上22:00后不得大声喧哗",
        "卫生间、厨房等公共区域使用后须及时清理",
        "严禁乙方饲养宠物",
        "乙方不得带领超过2名外来人员同时留宿"
    ],
    "penalties": [
        "乙方逾期支付租金，每逾期一天按月租金的3%收取滞纳金",
        "乙方提前退租，押金不予退还",
        "乙方损坏房屋设施或造成他人财产损失，由乙方承担全部责任",
        "乙方违反作息规定超过3次，甲方有权解除合同并不退还押金"
    ],
    "other_terms": [
        "本房屋已按规定进行改造并符合相关安全标准（如遇政府部门检查，乙方须配合）",
        "押金条款为合同重要条款，乙方确认已充分了解",
        "本房屋租金为不含税价格",
        "如房屋被政府认定为违建强制拆除，甲方不承担违约责任，仅退还剩余租金"
    ],
    "sign_date": "2024年7月25日",
    "sign_location": "杭州市西湖区"
})

# ========== 合同3：刚毕业大学生第一次租房（正常版） ==========
create_contract("03_毕业生标准租房合同.docx", {
    "title": "房屋租赁合同",
    "contract_no": "XS-2024-008",
    "lessor": "陈美玲（身份证：440106198608151234）",
    "lessee": "王鹏飞（身份证：410102199806013456）",
    "property_address": "广州市天河区天河路385号太古汇公寓B座1503室",
    "area": "45",
    "layout": "一室一厅一厨一卫",
    "condition": "精装修，配备格力空调、美的热水器、海尔冰箱、小天鹅洗衣机、床、沙发、衣柜、书桌等，拎包入住",
    "start_date": "2024年7月15日",
    "end_date": "2025年7月14日",
    "duration": "12个月",
    "monthly_rent": "3,800",
    "payment_method": "每月15日前以微信/支付宝转账方式支付",
    "deposit": "7,600（2个月租金）",
    "deposit_months": "2",
    "deposit_conditions": "租赁期满，乙方将房屋恢复原状（正常使用磨损除外）并结清全部费用后，甲方在3个工作日内无息退还押金",
    "water_fee": "按自来水公司标准由业主代收，每两个月结算",
    "electricity_fee": "持电卡至供电局充值，乙方自理",
    "gas_fee": "持燃气卡至燃气公司充值，乙方自理",
    "property_fee": "每月280元（含小区安保、公共绿化、卫生清洁）",
    "internet_fee": "已开通100M电信宽带，每月80元，乙方自理",
    "agency_fee": "业主自租，无中介费",
    "obligations": [
        "甲方应保证房屋权属清晰，交付时房屋及设施完好",
        "乙方应合理使用房屋及设施，发现损坏及时告知甲方",
        "租赁期间，乙方不得擅自将房屋转租、转借他人",
        "乙方须按时缴纳各项费用，不得以未入住等理由拖欠租金"
    ],
    "penalties": [
        "乙方逾期支付租金超过15日的，甲方有权要求乙方按日支付滞纳金（按银行同期贷款利率计算）",
        "乙方提前退租须提前30天书面通知甲方，并协助甲方带客看房",
        "乙方擅自转租的，甲方有权解除合同并要求乙方支付一个月租金的违约金"
    ],
    "other_terms": [
        "本合同自双方签字之日起生效",
        "合同一式两份，甲乙双方各执一份",
        "合同期满，乙方如需续租应提前一个月与甲方协商",
        "双方约定其他事项：本房屋仅供乙方本人居住使用"
    ],
    "sign_date": "2024年7月10日",
    "sign_location": "广州市天河区"
})

# ========== 合同4：虚假宣传/贷欺风险版 ==========
create_contract("04_租金贷风险合同.docx", {
    "title": "房屋租赁合同",
    "contract_no": "JF-2024-022",
    "lessor": "北京自如寓科技有限公司（统一社会信用代码：91110108MA01WXYZ12）",
    "lessee": "刘思琪（身份证：130102199904051234）",
    "property_address": "北京市海淀区中关村大街1号院8号楼501室",
    "area": "28",
    "layout": "一室一卫（品牌公寓）",
    "condition": "精装修全配，智能门锁、空调、冰箱、洗衣机、独立卫浴、密码衣柜",
    "start_date": "2024年9月1日",
    "end_date": "2026年8月31日",
    "duration": "24个月",
    "monthly_rent": "4,200",
    'payment_method': '通过"应客"APP以银行租金分期贷款方式支付（乙方授权甲方办理）',
    "deposit": "4,200（1个月租金）",
    "deposit_months": "1",
    "deposit_conditions": "合同到期后且乙方无违约行为时退还",
    "water_fee": "每月固定50元（不分用水量）",
    "electricity_fee": "按国标电价通过APP充值",
    "gas_fee": "无燃气",
    "property_fee": "每月180元（含每周两次公区保洁、代收快递、安保服务）",
    "internet_fee": "已接入百兆宽带，免费使用",
    "agency_fee": "无（品牌公寓）",
    "obligations": [
        "乙方须年付租金或通过甲方合作金融机构办理租金分期",
        "如乙方选择租金分期，还款义务由乙方承担，与甲方无关",
        "乙方须保持房屋及设施完好，配合甲方定期免费消杀保洁",
        "乙方提前退租须提前60天申请，并支付两个月租金作为违约金"
    ],
    "penalties": [
        "如乙方租金分期还款逾期，影响乙方征信记录，由乙方自行负责",
        "乙方逾期支付租金超过7天，甲方有权要求乙方支付100元/次逾期手续费",
        "乙方提前解约且经甲方同意的，须一次性支付剩余租期租金的30%作为违约金，并扣除全部押金",
        "租金分期一旦生效，无论乙方是否入住或使用房屋，贷款本息均由乙方承担"
    ],
    "other_terms": [
        "特别约定：乙方已知晓并同意以租金分期方式支付租金，授权甲方代为向金融机构申请",
        "本租赁关系一经成立，乙方不得以工作变动、距离偏远等理由要求解除",
        "若乙方中途退租，甲方有权直接扣除押金冲抵违约金，不足部分由乙方补缴",
        "如乙方逾期还租导致征信受损，乙方承诺自行承担全部后果，与甲方及金融机构无关"
    ],
    "sign_date": "2024年8月20日",
    "sign_location": "北京市海淀区"
})

# ========== 合同5：假房东/诈骗风险版 ==========
create_contract("05_假房东风险合同.docx", {
    "title": "房屋租赁合同",
    "contract_no": "JD-2024-033",
    "lessor": "周志远（身份证：310101198806127890，已与房东签署托管协议）",
    "lessee": "赵文静（身份证：500101199705061234）",
    "property_address": "成都市锦江区春熙路太古里旁王府大厦B座1201室",
    "area": "50",
    "layout": "两室一厅一厨一卫",
    "condition": "简装，配备空调2台、热水器、燃气灶、床2张、沙发、茶几、衣柜",
    "start_date": "2024年10月1日",
    "end_date": "2025年9月30日",
    "duration": "12个月",
    "monthly_rent": "2,200",
    "payment_method": "押一付三，首期租金及押金须在签约当日以现金方式支付",
    "deposit": "2,200（1个月租金）",
    "deposit_months": "1",
    "deposit_conditions": "合同到期归还房屋时退还",
    "water_fee": "预付500元，按实际使用量从押金中扣除",
    "electricity_fee": "预付300元，从押金中扣除",
    "gas_fee": "无",
    "property_fee": "每月120元",
    "internet_fee": "每月60元",
    "agency_fee": "一次性收取500元（服务费，含一次免费保洁）",
    "obligations": [
        "乙方须在签约当日支付首期租金、押金及各项预付款共计12,600元",
        "实际房东已全权委托本公司处理出租事宜，乙方无需联系原房东",
        "乙方须配合每季度一次的例行房屋检查（提前24小时通知）",
        "乙方不得在屋内进行任何形式的聚会或商业活动"
    ],
    "penalties": [
        "乙方逾期支付租金超过5日，视为自动退租，甲方有权立即收回房屋且押金不予退还",
        "如乙方提前退租，须提前45天书面通知，并支付两个月租金作为违约金",
        "乙方损坏房屋设施的，须按市场价的1.5倍赔偿"
    ],
    "other_terms": [
        "特别优惠：签约即送品牌床上用品一套（已包含在租金中）",
        "租金价格为此特别优惠价格，乙方不得向任何人透露实际成交价",
        "本合同签署后租金不予减免（工作变动、疾病等均不构成减免理由）",
        "如乙方欠费超过一个月，甲方有权断水断电且不构成违约",
        "争议解决：提交甲方所在地仲裁委员会仲裁（一裁终局）"
    ],
    "sign_date": "2024年9月28日",
    "sign_location": "成都市锦江区"
})

# ========== 合同6：实习生/短租版 ==========
create_contract("06_短租实习合同.docx", {
    "title": "房屋租赁合同（短租版）",
    "contract_no": "SX-2024-006",
    "lessor": "上海青租公寓管理有限公司（统一社会信用代码：91310000MA1FABC123）",
    "lessee": "陈伟（身份证：341201199903211234，实习生）",
    "property_address": "上海市浦东新区张江高科技园区碧波路690号张江微电子港3号楼203室",
    "area": "20",
    "layout": "单间（与其他房间共用厨卫，共4室）",
    "condition": "精装修，配备床、衣柜、书桌、空调、热水器，可使用公共厨房和客厅",
    "start_date": "2024年6月1日",
    "end_date": "2024年8月31日",
    "duration": "3个月（短租）",
    "monthly_rent": "2,100",
    "payment_method": "一次性付清全部租金及押金，以银行转账方式支付",
    "deposit": "2,100（1个月租金）",
    "deposit_months": "1",
    "deposit_conditions": "租赁期满乙方搬离且无损坏，房屋验收通过后7个工作日内无息退还",
    "water_fee": "已包含在租金中",
    "electricity_fee": "预付200元，退租时按实际用量结算，多退少补",
    "gas_fee": "无",
    "property_fee": "已包含在租金中",
    "internet_fee": "已包含在租金中（50M宽带）",
    "agency_fee": "无（企业合作公寓）",
    "obligations": [
        "乙方须为应届毕业实习生，租赁目的仅为工作期间居住",
        "乙方须提供实习证明或就业三方协议复印件",
        "公共区域（厨房、卫生间、客厅）使用后须及时清洁",
        "晚上23:00后须保持安静，不得影响邻居"
    ],
    "penalties": [
        "乙方逾期支付任何费用，每逾期一天按应付款项的1%收取滞纳金",
        "乙方提前退租须提前15天书面通知，押金退还50%；提前7天通知，押金退还25%；7天内不退押金",
        "乙方损坏房屋设施须照价赔偿"
    ],
    "other_terms": [
        "本合同专为短期实习学生设计，不支持续租",
        "租赁期满后，如乙方转为正式员工，可申请换签长租合同",
        "乙方离职或实习结束后须按时搬离，逾期按日租金的200%收取占用费",
        "合同一式两份，甲乙双方各执一份"
    ],
    "sign_date": "2024年5月25日",
    "sign_location": "上海市浦东新区"
})

# ========== 合同7：二房东转租风险版 ==========
create_contract("07_二房东转租风险合同.docx", {
    "title": "房屋租赁合同（转租）",
    "contract_no": "ES-2024-019",
    "lessor": "马超（身份证：420117199204057890，原租客，现转租人）",
    "lessee": "孙悦（身份证：220101199801011234）",
    "property_address": "武汉市洪山区光谷大道100号光谷国际广场A座1605室（原合同地址）",
    "area": "25（主卧带独立卫生间）",
    "layout": "三室一厅改四室，乙方承租其中一间主卧",
    "condition": "原装修，配备空调、热水器、床、衣柜、书桌",
    "start_date": "2024年8月1日",
    "end_date": "2025年1月31日",
    "duration": "6个月（仅限原合同剩余租期）",
    "monthly_rent": "1,800",
    "payment_method": "每月5日前以微信转账方式支付",
    "deposit": "1,800（1个月租金）",
    "deposit_months": "1",
    "deposit_conditions": "合同到期乙方搬离且无损坏后退还",
    "water_fee": "每月固定50元",
    "electricity_fee": "与室友按房间分摊（每间独立电表）",
    "gas_fee": "与室友共用，按人头分摊",
    "property_fee": "每月100元",
    "internet_fee": "每月45元",
    "agency_fee": "无",
    "obligations": [
        "甲方声明已获原房东口头同意转租，如因此产生纠纷由甲方承担责任",
        "乙方须遵守原租赁合同的全部条款",
        "乙方不得再将房屋转租、分租给第三方",
        "如原房东要求收回房屋，乙方须配合，押金由甲方向原房东追讨后返还乙方"
    ],
    "penalties": [
        "乙方逾期支付租金，每逾期一天按月租金的2%收取滞纳金",
        "乙方提前退租须提前20天通知，押金退还50%",
        "如原合同提前终止，本合同自动解除，甲方退还剩余租金，乙方不得索赔"
    ],
    "other_terms": [
        "重要提示：乙方确认已知晓本房屋为转租，存在被原房东收回的风险",
        "如遇国家或地方政策调整导致房屋需腾退，双方互不承担违约责任",
        "本合同效力依附于甲方与原房东的租赁合同，甲方须保证原合同至少持续到本合同到期日",
        "争议解决：协商解决，协商不成提交武汉市洪山区人民法院管辖"
    ],
    "sign_date": "2024年7月30日",
    "sign_location": "武汉市洪山区"
})

print("\n✅ 所有大学生租房合同样本已生成！")
print(f"📁 文件位置：D:\\aaa\\contract-review-copilot\\sample_contracts\\")
