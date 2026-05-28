from __future__ import annotations

from io import BytesIO
from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

REPORT_FONT_NAME = "Microsoft YaHei"
ZERO_WIDTH_CHAR_PATTERN = re.compile(r"[\u200b\u200c\u200d\u2060\ufeff]")
ITALIC_MARK_PATTERN = re.compile(r"(?<!\*)\*(.+?)\*(?!\*)")
MULTISPACE_PATTERN = re.compile(r"[ \t]{2,}")
LEADING_EMOJI_REPLACEMENTS = {
    "🔴": "[高危] ",
    "🟠": "[高风险] ",
    "🟡": "[提示] ",
    "🟢": "[提示] ",
    "⚪": "[提示] ",
    "⚠️": "[提示] ",
    "⚠": "[提示] ",
    "📋": "[说明] ",
    "✅": "[完成] ",
}
LEADING_LABEL_NORMALIZATIONS = (
    (re.compile(r"^\[说明\]\s*说明[:：]\s*"), "[说明] "),
    (re.compile(r"^\[提示\]\s*提示[:：]\s*"), "[提示] "),
)


def _clean_filename(source_filename: str | None) -> str:
    stem = Path(source_filename or "").stem.strip()
    return stem or "避坑指南"


def _apply_run_font(run, *, bold: bool | None = None, size: int | None = None) -> None:
    if bold is not None:
        run.bold = bold
    if size is not None:
        run.font.size = Pt(size)
    run.font.name = REPORT_FONT_NAME
    run._element.rPr.rFonts.set(qn("w:ascii"), REPORT_FONT_NAME)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), REPORT_FONT_NAME)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), REPORT_FONT_NAME)


def _normalize_inline_text(text: str) -> str:
    normalized = text.replace("\r\n", "\n").replace("\r", "\n").replace("\xa0", " ")
    for marker, replacement in LEADING_EMOJI_REPLACEMENTS.items():
        normalized = normalized.replace(marker, replacement)
    normalized = ZERO_WIDTH_CHAR_PATTERN.sub("", normalized).replace("\ufe0f", "")
    normalized = normalized.replace("**", "")
    normalized = normalized.replace("__", "")
    normalized = normalized.replace("`", "")
    normalized = ITALIC_MARK_PATTERN.sub(r"\1", normalized)
    normalized = MULTISPACE_PATTERN.sub(" ", normalized)
    for pattern, replacement in LEADING_LABEL_NORMALIZATIONS:
        normalized = pattern.sub(replacement, normalized)
    return normalized.strip()


def _iter_text_runs(text: str) -> list[tuple[str, bool]]:
    chunks: list[tuple[str, bool]] = []
    for part in re.split(r"(\*\*.*?\*\*)", text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) >= 4:
            content = _normalize_inline_text(part[2:-2])
            if content:
                chunks.append((content, True))
            continue
        content = _normalize_inline_text(part)
        if content:
            chunks.append((content, False))
    return chunks


def _append_text_runs(paragraph, text: str) -> None:
    for content, bold in _iter_text_runs(text):
        run = paragraph.add_run(content)
        _apply_run_font(run, bold=bold, size=11)


def _configure_document_styles(document: Document) -> None:
    styles_to_update = ("Normal", "Title", "Heading 1", "Heading 2")
    for style_name in styles_to_update:
        style = document.styles[style_name]
        style.font.name = REPORT_FONT_NAME
        style.font.size = Pt(11 if style_name == "Normal" else 12)
        style.element.rPr.rFonts.set(qn("w:ascii"), REPORT_FONT_NAME)
        style.element.rPr.rFonts.set(qn("w:hAnsi"), REPORT_FONT_NAME)
        style.element.rPr.rFonts.set(qn("w:eastAsia"), REPORT_FONT_NAME)


def _add_body_line(document: Document, line: str) -> None:
    if line == "---":
        return

    bullet_prefixes = ("- ", "• ", "· ")
    if line.startswith(bullet_prefixes):
        paragraph = document.add_paragraph()
        run = paragraph.add_run("• ")
        _apply_run_font(run, size=11)
        _append_text_runs(paragraph, line[2:].strip())
        return

    paragraph = document.add_paragraph()
    _append_text_runs(paragraph, line)


def build_report_docx(report_paragraphs: list[str], source_filename: str | None = None) -> bytes:
    document = Document()
    _configure_document_styles(document)

    title = document.add_heading("合同避坑指南", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        _apply_run_font(run, bold=True, size=18)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(f"来源文件：{_clean_filename(source_filename)}")
    _apply_run_font(subtitle_run, size=10)

    for paragraph in report_paragraphs:
        content = paragraph.strip()
        if not content:
            continue

        current_heading_level: int | None = None
        for raw_line in content.splitlines():
            line = _normalize_inline_text(raw_line)
            if not line:
                continue

            if line.startswith("## "):
                heading_text = line.replace("## ", "", 1).strip()
                if heading_text == "合同避坑指南":
                    continue
                heading = document.add_heading(heading_text, level=1)
                for run in heading.runs:
                    _apply_run_font(run, bold=True, size=14)
                current_heading_level = 1
                continue

            if line.startswith("### "):
                heading = document.add_heading(line.replace("### ", "", 1).strip(), level=2)
                for run in heading.runs:
                    _apply_run_font(run, bold=True, size=12)
                current_heading_level = 2
                continue

            if current_heading_level == 1 and line.startswith("##"):
                continue

            _add_body_line(document, line)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def build_report_download_name(source_filename: str | None = None) -> str:
    return f"{_clean_filename(source_filename)}_避坑指南.docx"
