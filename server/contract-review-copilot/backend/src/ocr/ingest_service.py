"""
Unified contract ingestion service for txt, docx, images, and PDFs.
"""
from __future__ import annotations

from dataclasses import asdict, dataclass
from io import BytesIO
from pathlib import Path

from docx import Document
from ..config import get_settings
from ..llm_client import extract_text_from_image

SUPPORTED_IMAGE_EXTENSIONS = {"jpg", "jpeg", "png", "webp"}
SUPPORTED_TEXT_EXTENSIONS = {"txt", "docx", "pdf"}
SUPPORTED_EXTENSIONS = SUPPORTED_TEXT_EXTENSIONS | SUPPORTED_IMAGE_EXTENSIONS


@dataclass(frozen=True)
class UploadedContractFile:
    filename: str
    content: bytes
    content_type: str | None = None

    @property
    def extension(self) -> str:
        return Path(self.filename).suffix.lower().removeprefix(".")


@dataclass(frozen=True)
class IngestedPageResult:
    page_index: int
    filename: str
    text: str
    average_confidence: float | None
    warnings: list[str]


@dataclass(frozen=True)
class ContractIngestResult:
    source_type: str
    display_name: str
    used_ocr_model: str | None
    merged_text: str
    pages: list[IngestedPageResult]
    warnings: list[str]

    def to_dict(self) -> dict:
        return {
            "source_type": self.source_type,
            "display_name": self.display_name,
            "used_ocr_model": self.used_ocr_model,
            "merged_text": self.merged_text,
            "pages": [asdict(page) for page in self.pages],
            "warnings": self.warnings,
        }


def _format_file_size(limit_bytes: int) -> str:
    megabytes = limit_bytes / (1024 * 1024)
    if float(megabytes).is_integer():
        return f"{int(megabytes)} MB"
    return f"{megabytes:.1f} MB"


def _format_megapixels(pixel_count: int) -> str:
    megapixels = pixel_count / 1_000_000
    if float(megapixels).is_integer():
        return f"{int(megapixels)} MP"
    return f"{megapixels:.1f} MP"


def _decode_text_file(file: UploadedContractFile) -> str:
    for encoding in ("utf-8", "utf-8-sig", "gb18030", "gbk"):
        try:
            return file.content.decode(encoding).strip()
        except UnicodeDecodeError:
            continue
    raise ValueError(f"{file.filename} 无法按文本文件读取，请检查编码格式。")


def _extract_docx_text(file: UploadedContractFile) -> str:
    document = Document(BytesIO(file.content))
    chunks: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            chunks.append(text)

    for table in document.tables:
        for row in table.rows:
            row_values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_values:
                chunks.append(" | ".join(row_values))

    return "\n".join(chunks).strip()


def _build_display_name(files: list[UploadedContractFile]) -> str:
    if len(files) == 1:
        return files[0].filename
    return f"{Path(files[0].filename).stem} 等 {len(files)} 页图片"


def _extract_text_from_uploaded_image(file: UploadedContractFile) -> tuple[str, str]:
    return extract_text_from_image(
        image_bytes=file.content,
        mime_type=file.content_type,
        filename=file.filename,
    )


def _read_image_size(file: UploadedContractFile) -> tuple[int, int] | None:
    try:
        from PIL import Image
    except ModuleNotFoundError:  # pragma: no cover - optional dependency
        return None

    try:
        with Image.open(BytesIO(file.content)) as image:
            return image.size
    except Exception:  # pragma: no cover - best-effort validation
        return None


def _count_pdf_pages(file: UploadedContractFile) -> int:
    import pypdfium2 as pdfium

    pdf = pdfium.PdfDocument(BytesIO(file.content))
    try:
        return len(pdf)
    finally:
        close = getattr(pdf, "close", None)
        if callable(close):
            close()


def _validate_file_size(file: UploadedContractFile) -> None:
    if file.extension in SUPPORTED_IMAGE_EXTENSIONS:
        return

    limit_bytes = get_settings().ocr_max_upload_file_bytes
    if len(file.content) > limit_bytes:
        raise ValueError(
            f"{file.filename} 超过单文件上传限制，请压缩到 {_format_file_size(limit_bytes)} 以内后重试。"
        )


def _validate_image_pixels(file: UploadedContractFile) -> None:
    max_pixels = get_settings().ocr_max_image_pixels
    if max_pixels <= 0:
        return

    image_size = _read_image_size(file)
    if image_size is None:
        return

    width, height = image_size
    pixel_count = width * height
    if pixel_count > max_pixels:
        raise ValueError(
            f"{file.filename} 图片分辨率过大，请控制在 {_format_megapixels(max_pixels)} 以内后重试。"
        )


def _validate_pdf_page_count(file: UploadedContractFile) -> None:
    max_pages = get_settings().ocr_max_pdf_pages
    page_count = _count_pdf_pages(file)
    if page_count > max_pages:
        raise ValueError(
            f"{file.filename} 共 {page_count} 页，超过 {max_pages} 页上限，请拆分后再上传。"
        )


def _validate_upload_constraints(files: list[UploadedContractFile]) -> None:
    if not files:
        raise ValueError("请至少上传一个文件。")

    for file in files:
        if not file.content:
            raise ValueError(f"{file.filename or '上传文件'} 内容为空。")
        if file.extension not in SUPPORTED_EXTENSIONS:
            raise ValueError(
                "仅支持 TXT、DOCX、PDF 以及 JPG/PNG/WEBP 合同材料。"
            )
        _validate_file_size(file)

    all_images = all(file.extension in SUPPORTED_IMAGE_EXTENSIONS for file in files)
    if len(files) > 1 and not all_images:
        raise ValueError(
            "一次仅支持批量上传多张图片，TXT、DOCX、PDF 请单独上传。"
        )

    if all_images:
        max_batch_images = get_settings().ocr_max_batch_images
        if len(files) > max_batch_images:
            raise ValueError(
                f"一次最多上传 {max_batch_images} 张合同图片，请分批上传。"
            )
        for file in files:
            _validate_image_pixels(file)
        return

    file = files[0]
    if file.extension == "pdf":
        _validate_pdf_page_count(file)


def _ingest_image_files(
    files: list[UploadedContractFile],
    *,
    display_name: str,
    source_type: str,
) -> ContractIngestResult:
    pages: list[IngestedPageResult] = []
    warnings: list[str] = []
    used_model: str | None = None

    for page_index, file in enumerate(files, start=1):
        try:
            page_text, page_model = _extract_text_from_uploaded_image(file)
        except Exception as exc:
            warnings.append(f"第 {page_index} 页 OCR 失败：{exc}")
            continue

        used_model = page_model
        pages.append(
            IngestedPageResult(
                page_index=page_index,
                filename=file.filename,
                text=page_text,
                average_confidence=None,
                warnings=[],
            )
        )

    merged_text = "\n\n".join(page.text.strip() for page in pages if page.text.strip()).strip()
    if not merged_text:
        if warnings:
            raise RuntimeError("；".join(warnings))
        raise RuntimeError("未能从上传材料中提取出有效合同文本。")

    return ContractIngestResult(
        source_type=source_type,
        display_name=display_name,
        used_ocr_model=used_model,
        merged_text=merged_text,
        pages=pages,
        warnings=warnings,
    )


def _render_pdf_to_images(file: UploadedContractFile) -> list[UploadedContractFile]:
    import pypdfium2 as pdfium

    pdf = pdfium.PdfDocument(BytesIO(file.content))
    stem = Path(file.filename).stem
    rendered_files: list[UploadedContractFile] = []

    try:
        for page_index in range(len(pdf)):
            page = pdf[page_index]
            bitmap = page.render(scale=2.2)
            pil_image = bitmap.to_pil()
            buffer = BytesIO()
            pil_image.save(buffer, format="PNG")
            rendered_files.append(
                UploadedContractFile(
                    filename=f"{stem}-page-{page_index + 1}.png",
                    content=buffer.getvalue(),
                    content_type="image/png",
                )
            )
    finally:
        close = getattr(pdf, "close", None)
        if callable(close):
            close()

    return rendered_files


def validate_contract_uploads(files: list[UploadedContractFile]) -> None:
    _validate_upload_constraints(files)


def ingest_contract_files(files: list[UploadedContractFile]) -> ContractIngestResult:
    validate_contract_uploads(files)

    all_images = all(file.extension in SUPPORTED_IMAGE_EXTENSIONS for file in files)
    if all_images:
        return _ingest_image_files(
            files,
            display_name=_build_display_name(files),
            source_type="image_batch",
        )

    file = files[0]
    extension = file.extension
    if extension == "txt":
        text = _decode_text_file(file)
        if not text:
            raise ValueError("TXT 文件内容为空。")
        return ContractIngestResult(
            source_type="txt",
            display_name=file.filename,
            used_ocr_model=None,
            merged_text=text,
            pages=[
                IngestedPageResult(
                    page_index=1,
                    filename=file.filename,
                    text=text,
                    average_confidence=None,
                    warnings=[],
                )
            ],
            warnings=[],
        )

    if extension == "docx":
        text = _extract_docx_text(file)
        if not text:
            raise ValueError("DOCX 文件未提取到有效文本。")
        return ContractIngestResult(
            source_type="docx",
            display_name=file.filename,
            used_ocr_model=None,
            merged_text=text,
            pages=[
                IngestedPageResult(
                    page_index=1,
                    filename=file.filename,
                    text=text,
                    average_confidence=None,
                    warnings=[],
                )
            ],
            warnings=[],
        )

    if extension == "pdf":
        rendered_files = _render_pdf_to_images(file)
        if not rendered_files:
            raise RuntimeError("PDF 未能渲染出可识别页面。")
        return _ingest_image_files(rendered_files, display_name=file.filename, source_type="pdf_ocr")

    raise ValueError("仅支持 TXT、DOCX、PDF 以及 JPG/PNG/WEBP 合同材料。")
