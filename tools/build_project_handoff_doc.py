from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/wsr/agent/ContractReviewCopilot-iOS")
OUT = ROOT / "CTSafe_iOS_项目交接与后续开发计划.docx"


COLORS = {
    "blue": "2E74B5",
    "dark_blue": "1F4D78",
    "muted": "666666",
    "fill": "F2F4F7",
    "green": "137333",
    "orange": "B06000",
    "red": "B3261E",
}


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(10)


def set_cell_width(cell, width_inches: float) -> None:
    cell.width = Inches(width_inches)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def style_table(table, widths: list[float] | None = None) -> None:
    table.autofit = False
    table.style = "Table Grid"
    for row_index, row in enumerate(table.rows):
        for idx, cell in enumerate(row.cells):
            if widths and idx < len(widths):
                set_cell_width(cell, widths[idx])
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                    run.font.size = Pt(10)
            if row_index == 0:
                shade_cell(cell, COLORS["fill"])
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True


def add_heading(doc: Document, text: str, level: int = 1):
    paragraph = doc.add_heading(text, level=level)
    for run in paragraph.runs:
        run.font.name = "Calibri"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        if level == 1:
            run.font.color.rgb = RGBColor.from_string(COLORS["blue"])
        else:
            run.font.color.rgb = RGBColor.from_string(COLORS["dark_blue"])
    return paragraph


def add_para(doc: Document, text: str = "", *, bold_prefix: str | None = None):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.1
    if bold_prefix and text.startswith(bold_prefix):
        prefix = paragraph.add_run(bold_prefix)
        prefix.bold = True
        body = paragraph.add_run(text[len(bold_prefix) :])
        runs = [prefix, body]
    else:
        runs = [paragraph.add_run(text)]
    for run in runs:
        run.font.name = "Calibri"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(11)
    return paragraph


def add_bullet(doc: Document, text: str):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(11)
    return paragraph


def add_number(doc: Document, text: str):
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(11)
    return paragraph


def add_status_table(doc: Document):
    rows = [
        ("服务器连接", "已完成", "Mac SSH key 已生成并配置到服务器，别名 myserver 可用。未在本文档记录任何密码。"),
        ("服务端扫描", "已完成", "确认 Web/FastAPI 项目位于服务器 /root/contract-review-copilot，本地镜像位于 server/contract-review-copilot。"),
        ("后端同步层", "已完成", "用户文档、审查会话、风险项、报告、聊天消息持久化模型和查询接口均已实现。"),
        ("Web 端迁移", "已完成", "历史、报告、问答已读取云端接口，localStorage 保留兼容显示。"),
        ("iOS 工程", "已完成", "SwiftUI 工程可在 iPhone 17 Simulator 编译、安装、启动。14 个 Swift 文件共 2,587 行代码。"),
        ("登录/注册", "已完成", "API 拼接问题已修复，注册 UI 与验证码流程已实现；OAuth 深度链接已接入。"),
        ("Doge UI 主题", "已完成", "DogeTheme 设计系统、PixelPanel、DogeBadge、RiskPill 已全部实现，像素风贯穿所有页面。"),
        ("文件上传与审查流", "已完成", "multipart 上传、OCR 文本提取、SSE 审查流（6 阶段进度条）、断点确认均已实现。"),
        ("合同审查报告", "已完成", "风险总览、高/中/低风险计数、可展开问题行、RiskPill、状态徽章均已实现。"),
        ("问答对话", "已完成", "会话选择器、打字指示器、关键发现横条、聊天消息收发均已实现。"),
        ("历史记录", "已完成", "搜索（文件名+关键词）、风险等级筛选、下拉刷新、统计数据栏均已实现。"),
    ]
    table = doc.add_table(rows=1, cols=3)
    headers = ["事项", "状态", "说明"]
    widths = [1.35, 1.35, 3.8]
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True)
    for item, status, note in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], item)
        set_cell_text(cells[1], status)
        set_cell_text(cells[2], note)
    style_table(table, widths)


def add_next_steps_table(doc: Document):
    rows = [
        ("P0", "真实账号端到端验证", "用真实 ctsafe.top 账号在模拟器完成登录、注册、验证码邮件全流程验证。"),
        ("P0", "跨端同步验收", "Web 上传合同后 iOS 历史可见；iOS 发起问答后 Web 能看到同一会话聊天记录。"),
        ("P0", "ReportView 补齐", "「查看原文」按钮当前为空操作，需实现原文展示功能。"),
        ("P1", "OAuth 真机验证", "Google/GitHub OAuth 已接入 ctsafe://auth 回调，需在真机或模拟器完整跑通登录流程。"),
        ("P1", "错误处理与离线优化", "网络错误目前仅显示重试按钮，需增加离线缓存、自动重试、错误分类等鲁棒性。"),
        ("P1", "服务端测试补充", "补用户隔离、审查落库、聊天跨端可读、文档归属等自动化测试。"),
        ("P1", "Web 端回归", "重新测试上传、审查、历史、报告、问答，确认迁移云端数据后没有破坏原流程。"),
        ("P2", "代码质量提升", "提取重复的 fileIcon()/formatDate() 到公共模块；将 View 直接调用 API 改为 ViewModel 模式；添加单元测试。"),
        ("P2", "国际化与本地化", "当前所有 UI 字符串为硬编码中文，需建立 Localizable.strings 支持中英双语。"),
        ("P2", "真机与发布准备", "配置 Apple Developer Team、Bundle ID、图标、权限文案、生产构建和 TestFlight。"),
    ]
    table = doc.add_table(rows=1, cols=3)
    headers = ["优先级", "任务", "验收标准"]
    widths = [0.75, 1.75, 4.0]
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True)
    for priority, task, criteria in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], priority)
        set_cell_text(cells[1], task)
        set_cell_text(cells[2], criteria)
    style_table(table, widths)


def build_doc() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run("CTSafe iOS + 云端同步项目交接与后续开发计划")
    run.bold = True
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor.from_string(COLORS["blue"])

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(18)
    srun = subtitle.add_run("整理日期：2026-05-28    工作目录：/Users/wsr/agent/ContractReviewCopilot-iOS")
    srun.font.name = "Calibri"
    srun._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    srun.font.size = Pt(10)
    srun.font.color.rgb = RGBColor.from_string(COLORS["muted"])

    add_heading(doc, "1. 项目目标", 1)
    add_para(
        doc,
        "目标是从现有 ctsafe Web/FastAPI 项目出发，增加用户资产云端同步层，并创建 SwiftUI iOS 客户端。第一版要求同一账号在 Web 与 iOS 上共享合同、审查报告、问答和历史记录。",
    )
    add_para(
        doc,
        "当前同步定义：服务端同源数据 + 页面进入/刷新时拉取最新数据 + 审查和聊天继续使用 SSE；暂不做 APNs 推送、离线审查、团队协作和旧本地历史批量回填。",
    )

    add_heading(doc, "2. 已完成事项总览", 1)
    add_status_table(doc)

    add_heading(doc, "3. 服务器与连接配置", 1)
    add_bullet(doc, "已为本机 Mac 生成 SSH key，并把公钥加入服务器 root 用户授权。")
    add_bullet(doc, "已配置 ~/.ssh/config，Host 别名为 myserver，可通过 ssh myserver 连接。")
    add_bullet(doc, "服务器地址为 107.174.71.175，主机名验证为 racknerd-7d6c0a9。")
    add_bullet(doc, "安全说明：用户曾提供服务器密码用于初次配置，但本文档不记录密码，也不记录私钥内容。")

    add_heading(doc, "4. 服务端与 Web 项目现状", 1)
    add_para(doc, "服务端主项目位于服务器 /root/contract-review-copilot，本地镜像位于 /Users/wsr/agent/ContractReviewCopilot-iOS/server/contract-review-copilot。")
    add_para(doc, "技术栈：React/Vite 前端、FastAPI 后端、Postgres + pgvector、Redis、Docker Compose。")
    add_bullet(doc, "新增后端持久化模块：backend/src/services/sync_store.py。")
    add_bullet(doc, "新增数据表：user_documents、review_sessions、review_findings、chat_messages。")
    add_bullet(doc, "新增用户隔离接口：GET /api/documents、GET /api/documents/{id}、GET /api/review-sessions、GET /api/review-sessions/{id}、GET /api/review-sessions/{id}/chat。")
    add_bullet(doc, "改造审查、聊天、OCR 相关流程：审查完成后写入报告和风险项；聊天接口写入消息；OCR 上传结果绑定用户文档。")
    add_bullet(doc, "Web 历史页、报告页、问答页已开始读取云端数据，旧 localStorage 只保留临时兼容。")
    add_bullet(doc, "已在服务器重新构建并启动 backend/frontend 容器，/health 和未登录 /api/documents 401 校验通过。")

    add_heading(doc, "5. iOS 工程现状", 1)
    add_para(doc, "iOS 工程路径：/Users/wsr/agent/ContractReviewCopilot-iOS/CTSafeiOS/CTSafeiOS.xcodeproj。共 14 个 Swift 文件，2,587 行代码，零第三方依赖。")
    add_para(doc, "技术栈：Swift 6.0 + SwiftUI + URLSession async/await + Security Keychain + SSE 自定义流式客户端。", bold_prefix="技术栈：")
    add_bullet(doc, "App 入口（CTSafeApp.swift）与认证门控（RootView.swift）：OAuth 深度链接 ctsafe://auth?token=... 已实现。")
    add_bullet(doc, "登录/注册（LoginView.swift, 319 行）：邮箱密码登录、注册验证码、密码策略校验、Google/GitHub OAuth 按钮均已实现。")
    add_bullet(doc, "首页（HomeView.swift, 462 行）：文件导入器（PDF/DOCX/DOC/TXT/JPG/PNG/WEBP）、multipart 上传、SSE 审查流（6 阶段进度条）、断点自动确认均已实现。")
    add_bullet(doc, "问答对话（ChatView.swift, 411 行）：会话选择器、打字指示器、关键发现横条、消息收发均已实现。")
    add_bullet(doc, "审查报告（ReportView.swift, 351 行）：风险总览、高/中/低计数卡片、可展开问题行、RiskPill、状态徽章均已实现。")
    add_bullet(doc, "历史记录（HistoryView.swift, 251 行）：搜索、风险筛选、下拉刷新、统计栏均已实现。")
    add_bullet(doc, "设置页（SettingsView.swift, 65 行）：用户信息展示、AI 免责声明、登出功能。")
    add_bullet(doc, "网络层（APIClient.swift, 120 行）：URL 路径拼接 bug 已修复，multipart 文件上传已实现。SSEClient.swift（103 行）支持 AsyncThrowingStream 流式解析。")
    add_bullet(doc, "认证与存储（AuthStore.swift 108 行 + KeychainStore.swift 39 行）：JWT Keychain 持久化、注册/登录/OAuth token 接受均已实现。")
    add_bullet(doc, "设计系统（Theme.swift, 61 行）：DogeTheme 颜色体系、PixelPanel 卡片修饰器、DogeBadge 吉祥物、RiskPill 风险标签均已实现。")

    add_heading(doc, "6. 已修复的关键问题", 1)
    add_para(doc, "登录 URL 拼接问题：APIClient 使用 baseURL=https://ctsafe.top/api，但调用 path 以 /auth/login 开头；Swift URL appending(path:) 在该情况下可能覆盖 /api 前缀。")
    add_bullet(doc, "已修复 APIClient.url(for:)：去掉 path 前导斜杠后再拼到 https://ctsafe.top/api。")
    add_bullet(doc, "已增强错误解析：同时读取 error 和 detail 字段，方便显示 FastAPI 401/422 等错误。")
    add_bullet(doc, "已新增 iOS 注册请求模型：SendCodeRequest、RegisterRequest、BasicAuthResponse。")
    add_bullet(doc, "已在 AuthStore 中新增 sendRegisterCode 与 register 方法。")
    add_bullet(doc, "SSE 流式客户端已支持 AsyncThrowingStream，自定义 URLSession 超时配置（300s 请求 / 600s 资源）适配长时间审查任务。")

    add_heading(doc, "7. Xcode 运行方式", 1)
    add_number(doc, "打开 /Users/wsr/agent/ContractReviewCopilot-iOS/CTSafeiOS/CTSafeiOS.xcodeproj。")
    add_number(doc, "顶部 Scheme 选择 CTSafeiOS。")
    add_number(doc, "运行设备选择 iPhone 17 或其他 iOS Simulator。")
    add_number(doc, "点击左上角 Run 三角按钮。")
    add_number(doc, "如果模拟器安装失败，先 Clean Build Folder 或卸载模拟器里的旧 com.ctsafe.ios，再重新 Run。")

    add_heading(doc, "8. 还需要做什么", 1)
    add_next_steps_table(doc)

    add_heading(doc, "9. 当前风险与注意事项", 1)
    add_bullet(doc, "登录/注册流程代码已完整，但尚未用真实账号端到端验证，尤其验证码邮件投递、token 写入 Keychain、/api/auth/me 初始化状态。")
    add_bullet(doc, "OAuth 移动端回调已设计为 ctsafe://auth?token=...，但 Google/GitHub 控制台回调配置和模拟器跳转仍需实际验证。")
    add_bullet(doc, "ReportView「查看原文」按钮为空操作，需补充实现。")
    add_bullet(doc, "Web 与 iOS 的同步基础接口已存在，但需要 E2E 验证同一账号跨端报告、历史和聊天一致。")
    add_bullet(doc, "无本地数据缓存：所有审查会话、聊天历史、文档均依赖网络实时拉取，离线不可用。")
    add_bullet(doc, "代码架构为扁平 MVVM-lite，View 直接调用 APIClient，无 ViewModel 层、无依赖注入、无单元测试。重复代码（fileIcon/formatDate）散布多个文件。")
    add_bullet(doc, "APIClient.shared 使用 nonisolated(unsafe) 单例，tokenProvider 闭包为可变状态，在 Swift 6 严格并发下存在隐患。")
    add_bullet(doc, "baseURL 硬编码为 https://ctsafe.top/api，无环境切换机制。")

    add_heading(doc, "10. 推荐下一步执行顺序", 1)
    add_number(doc, "用真实账号在模拟器完成登录、注册、验证码全流程验证。")
    add_number(doc, "Web 与 iOS 跨端同步验收：Web 上传合同后 iOS 历史可见，iOS 问答后 Web 可见。")
    add_number(doc, "补全 ReportView「查看原文」功能。")
    add_number(doc, "在真机或模拟器验证 Google/GitHub OAuth 完整流程。")
    add_number(doc, "提取重复代码、引入 ViewModel 层、添加单元测试。")
    add_number(doc, "配置 Apple Developer Team、Bundle ID、图标，准备 TestFlight 发布。")

    section = doc.add_section(WD_SECTION.CONTINUOUS)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    frun = footer.add_run("CTSafe iOS 项目交接文档")
    frun.font.size = Pt(9)
    frun.font.color.rgb = RGBColor.from_string(COLORS["muted"])

    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
    print(OUT)
